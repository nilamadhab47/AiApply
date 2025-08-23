"""
Template Application Service
Applies extracted resume data to DOCX templates using python-docx
"""

import os
import json
import logging
from typing import Dict, List, Any, Optional
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import io
from datetime import datetime
import re

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class TemplateData(dict):
    """Enhanced dictionary for template data with safe access"""
    
    def get_safe(self, key: str, default: str = "") -> str:
        """Get value with safe fallback"""
        value = self.get(key, default)
        return str(value) if value else default
    
    def get_list(self, key: str, default: List = None) -> List:
        """Get list value with safe fallback"""
        value = self.get(key, default or [])
        return value if isinstance(value, list) else default or []

class TemplateApplicator:
    """Applies extracted resume data to DOCX templates"""
    
    def __init__(self, templates_dir: str = "templates"):
        self.templates_dir = templates_dir
        self.ensure_templates_directory()
        
    def ensure_templates_directory(self):
        """Create templates directory if it doesn't exist"""
        if not os.path.exists(self.templates_dir):
            os.makedirs(self.templates_dir)
            logger.info(f"Created templates directory: {self.templates_dir}")
    
    def extract_resume_data(self, analysis: 'CompleteResumeAnalysis') -> TemplateData:
        """Extract and structure data from resume analysis for template application"""
        data = TemplateData()
        
        # Initialize with defaults
        data.update({
            'full_name': '',
            'email': '',
            'phone': '',
            'address': '',
            'linkedin': '',
            'summary': '',
            'experience': [],
            'education': [],
            'skills': [],
            'projects': [],
            'certifications': [],
            'languages': []
        })
        
        # Process each section
        for section in analysis.sections:
            section_type = section.type.lower()
            content = section.content.strip()
            
            if section_type == 'personal':
                self._extract_personal_info(content, data)
            elif section_type in ['summary', 'objective', 'profile']:
                data['summary'] = content
            elif section_type in ['experience', 'work']:
                data['experience'] = self._extract_experience(content)
            elif section_type == 'education':
                data['education'] = self._extract_education(content)
            elif section_type == 'skills':
                data['skills'] = self._extract_skills(content)
            elif section_type == 'projects':
                data['projects'] = self._extract_projects(content)
            elif section_type in ['certifications', 'certificates']:
                data['certifications'] = self._extract_certifications(content)
        
        # Add metadata
        data['ats_score'] = analysis.ats_analysis.overall_score
        data['generation_date'] = datetime.now().strftime("%B %Y")
        
        logger.info(f"Extracted data for: {data.get('full_name', 'Unknown')}")
        return data
    
    def _extract_personal_info(self, content: str, data: TemplateData):
        """Extract personal information from content"""
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Extract email
            email_match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', line)
            if email_match:
                data['email'] = email_match.group()
                continue
            
            # Extract phone
            phone_match = re.search(r'(\+?1?[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})', line)
            if phone_match:
                data['phone'] = phone_match.group()
                continue
            
            # Extract LinkedIn
            if 'linkedin.com' in line.lower():
                data['linkedin'] = line
                continue
            
            # If line doesn't contain contact info and we don't have a name yet, it's probably the name
            if not data.get('full_name') and not any(char in line for char in '@()+-0123456789'):
                if len(line.split()) >= 2:  # Likely a full name
                    data['full_name'] = line
                    continue
            
            # Remaining lines could be address
            if not data.get('address') and any(word in line.lower() for word in ['street', 'ave', 'road', 'city', 'state']):
                data['address'] = line
    
    def _extract_experience(self, content: str) -> List[Dict[str, str]]:
        """Extract work experience entries"""
        experiences = []
        
        # Split by common job separators
        entries = re.split(r'\n(?=[A-Z].*(?:20\d{2}|19\d{2}))', content)
        
        for entry in entries:
            entry = entry.strip()
            if not entry or len(entry) < 20:
                continue
            
            lines = entry.split('\n')
            exp_data = {
                'title': '',
                'company': '',
                'duration': '',
                'description': ''
            }
            
            # First line usually contains title and/or company
            first_line = lines[0].strip()
            
            # Look for dates in first few lines
            date_pattern = r'(20\d{2}|19\d{2})'
            date_line = ''
            for line in lines[:3]:
                if re.search(date_pattern, line):
                    date_line = line.strip()
                    exp_data['duration'] = date_line
                    break
            
            # Extract title and company from first line
            if ' at ' in first_line:
                parts = first_line.split(' at ', 1)
                exp_data['title'] = parts[0].strip()
                exp_data['company'] = parts[1].strip()
            elif ' - ' in first_line and not re.search(date_pattern, first_line):
                parts = first_line.split(' - ', 1)
                exp_data['title'] = parts[0].strip()
                exp_data['company'] = parts[1].strip()
            else:
                exp_data['title'] = first_line
                # Look for company in next lines
                for line in lines[1:3]:
                    if line.strip() and not re.search(date_pattern, line):
                        exp_data['company'] = line.strip()
                        break
            
            # Description is remaining content
            desc_lines = []
            for line in lines[1:]:
                if line.strip() and line.strip() != date_line and line.strip() != exp_data['company']:
                    desc_lines.append(line.strip())
            
            exp_data['description'] = '\n'.join(desc_lines)
            
            if exp_data['title'] or exp_data['company']:
                experiences.append(exp_data)
        
        return experiences[:5]  # Limit to 5 most recent
    
    def _extract_education(self, content: str) -> List[Dict[str, str]]:
        """Extract education entries"""
        education = []
        
        entries = content.split('\n\n')
        
        for entry in entries:
            entry = entry.strip()
            if not entry or len(entry) < 10:
                continue
            
            lines = [line.strip() for line in entry.split('\n') if line.strip()]
            
            edu_data = {
                'degree': '',
                'school': '',
                'year': '',
                'details': ''
            }
            
            # Look for degree keywords
            degree_keywords = ['bachelor', 'master', 'phd', 'doctorate', 'associate', 'certificate', 'diploma']
            
            for line in lines:
                if any(keyword in line.lower() for keyword in degree_keywords):
                    edu_data['degree'] = line
                elif re.search(r'20\d{2}|19\d{2}', line):
                    edu_data['year'] = line
                elif not edu_data['school'] and len(line) > 5:
                    edu_data['school'] = line
                else:
                    if edu_data['details']:
                        edu_data['details'] += '\n' + line
                    else:
                        edu_data['details'] = line
            
            if edu_data['degree'] or edu_data['school']:
                education.append(edu_data)
        
        return education
    
    def _extract_skills(self, content: str) -> List[str]:
        """Extract skills list"""
        skills = []
        
        # Split by common separators
        for separator in [',', '•', '·', '\n', ';']:
            if separator in content:
                skills = [skill.strip() for skill in content.split(separator)]
                break
        
        if not skills:
            # Fallback: split by spaces and take meaningful words
            words = content.split()
            skills = [word.strip('.,;') for word in words if len(word) > 2]
        
        # Clean and filter skills
        cleaned_skills = []
        for skill in skills:
            skill = skill.strip('.,;•·-')
            if skill and len(skill) > 1 and len(skill) < 50:
                cleaned_skills.append(skill)
        
        return cleaned_skills[:20]  # Limit to 20 skills
    
    def _extract_projects(self, content: str) -> List[Dict[str, str]]:
        """Extract project entries"""
        projects = []
        
        entries = content.split('\n\n')
        
        for entry in entries:
            entry = entry.strip()
            if not entry or len(entry) < 20:
                continue
            
            lines = entry.split('\n')
            project_data = {
                'name': lines[0].strip(),
                'description': '\n'.join(lines[1:]).strip()
            }
            
            projects.append(project_data)
        
        return projects[:3]  # Limit to 3 projects
    
    def _extract_certifications(self, content: str) -> List[str]:
        """Extract certifications list"""
        certs = []
        
        lines = content.split('\n')
        for line in lines:
            line = line.strip('•·-').strip()
            if line and len(line) > 5:
                certs.append(line)
        
        return certs[:5]  # Limit to 5 certifications
    
    def apply_template(self, template_name: str, resume_data: TemplateData) -> bytes:
        """Apply resume data to a specific template"""
        template_path = os.path.join(self.templates_dir, f"{template_name}.docx")
        
        if not os.path.exists(template_path):
            # Create a basic template if it doesn't exist
            logger.warning(f"Template {template_name} not found, creating basic template")
            return self._create_basic_resume(resume_data)
        
        try:
            # Load template document
            doc = Document(template_path)
            
            # Replace placeholders in document
            self._replace_placeholders(doc, resume_data)
            
            # Save to bytes
            doc_bytes = io.BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            
            logger.info(f"Successfully applied template {template_name}")
            return doc_bytes.getvalue()
            
        except Exception as e:
            logger.error(f"Error applying template {template_name}: {e}")
            # Fallback to basic resume
            return self._create_basic_resume(resume_data)
    
    def _replace_placeholders(self, doc: Document, data: TemplateData):
        """Replace placeholders in document with actual data"""
        
        # Define placeholder mappings
        placeholders = {
            '{{FULL_NAME}}': data.get_safe('full_name'),
            '{{EMAIL}}': data.get_safe('email'),
            '{{PHONE}}': data.get_safe('phone'),
            '{{ADDRESS}}': data.get_safe('address'),
            '{{LINKEDIN}}': data.get_safe('linkedin'),
            '{{SUMMARY}}': data.get_safe('summary'),
            '{{GENERATION_DATE}}': data.get_safe('generation_date'),
        }
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for placeholder, value in placeholders.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, value)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for placeholder, value in placeholders.items():
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, value)
        
        # Handle complex sections (experience, education, skills)
        self._replace_experience_section(doc, data)
        self._replace_education_section(doc, data)
        self._replace_skills_section(doc, data)
    
    def _replace_experience_section(self, doc: Document, data: TemplateData):
        """Replace experience section with actual experience data"""
        experience_data = data.get_list('experience')
        
        # Find experience placeholder
        for paragraph in doc.paragraphs:
            if '{{EXPERIENCE_SECTION}}' in paragraph.text:
                # Clear the placeholder
                paragraph.text = ''
                
                # Add experience entries
                for exp in experience_data:
                    # Job title and company
                    title_para = paragraph._element.getparent().insert_after(paragraph._element)
                    title_para = doc.paragraphs[len(doc.paragraphs)-1]
                    title_para.text = f"{exp['title']} - {exp['company']}"
                    title_para.style = 'Heading 3'
                    
                    # Duration
                    if exp['duration']:
                        duration_para = paragraph._element.getparent().insert_after(title_para._element)
                        duration_para = doc.paragraphs[len(doc.paragraphs)-1]
                        duration_para.text = exp['duration']
                    
                    # Description
                    if exp['description']:
                        desc_para = paragraph._element.getparent().insert_after(paragraph._element)
                        desc_para = doc.paragraphs[len(doc.paragraphs)-1]
                        desc_para.text = exp['description']
                
                break
    
    def _replace_education_section(self, doc: Document, data: TemplateData):
        """Replace education section with actual education data"""
        education_data = data.get_list('education')
        
        for paragraph in doc.paragraphs:
            if '{{EDUCATION_SECTION}}' in paragraph.text:
                paragraph.text = ''
                
                for edu in education_data:
                    edu_text = f"{edu['degree']} - {edu['school']}"
                    if edu['year']:
                        edu_text += f" ({edu['year']})"
                    
                    edu_para = paragraph._element.getparent().insert_after(paragraph._element)
                    edu_para = doc.paragraphs[len(doc.paragraphs)-1]
                    edu_para.text = edu_text
                
                break
    
    def _replace_skills_section(self, doc: Document, data: TemplateData):
        """Replace skills section with actual skills data"""
        skills_data = data.get_list('skills')
        
        for paragraph in doc.paragraphs:
            if '{{SKILLS_SECTION}}' in paragraph.text:
                skills_text = ', '.join(skills_data[:15])  # Limit skills display
                paragraph.text = skills_text
                break
    
    def _create_basic_resume(self, data: TemplateData) -> bytes:
        """Create a basic resume when no template is available"""
        doc = Document()
        
        # Header
        header = doc.add_heading(data.get_safe('full_name'), 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact info
        contact_info = []
        if data.get_safe('email'):
            contact_info.append(data.get_safe('email'))
        if data.get_safe('phone'):
            contact_info.append(data.get_safe('phone'))
        if data.get_safe('address'):
            contact_info.append(data.get_safe('address'))
        
        if contact_info:
            contact_para = doc.add_paragraph(' | '.join(contact_info))
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Summary
        if data.get_safe('summary'):
            doc.add_heading('Professional Summary', level=1)
            doc.add_paragraph(data.get_safe('summary'))
        
        # Experience
        experience_data = data.get_list('experience')
        if experience_data:
            doc.add_heading('Work Experience', level=1)
            for exp in experience_data:
                exp_header = f"{exp['title']} - {exp['company']}"
                if exp['duration']:
                    exp_header += f" ({exp['duration']})"
                doc.add_heading(exp_header, level=2)
                if exp['description']:
                    doc.add_paragraph(exp['description'])
        
        # Education
        education_data = data.get_list('education')
        if education_data:
            doc.add_heading('Education', level=1)
            for edu in education_data:
                edu_text = f"{edu['degree']} - {edu['school']}"
                if edu['year']:
                    edu_text += f" ({edu['year']})"
                doc.add_paragraph(edu_text)
        
        # Skills
        skills_data = data.get_list('skills')
        if skills_data:
            doc.add_heading('Skills', level=1)
            doc.add_paragraph(', '.join(skills_data[:15]))
        
        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        logger.info("Created basic resume template")
        return doc_bytes.getvalue()
    
    def get_available_templates(self) -> List[str]:
        """Get list of available template names"""
        if not os.path.exists(self.templates_dir):
            return []
        
        templates = []
        for file in os.listdir(self.templates_dir):
            if file.endswith('.docx'):
                templates.append(file.replace('.docx', ''))
        
        return templates
    
    def create_sample_templates(self):
        """Create sample templates for testing"""
        logger.info("Creating sample templates...")
        
        # Professional Template
        self._create_professional_template()
        
        # Modern Template  
        self._create_modern_template()
        
        logger.info("Sample templates created successfully")
    
    def _create_professional_template(self):
        """Create a professional template with placeholders"""
        doc = Document()
        
        # Header
        header = doc.add_heading('{{FULL_NAME}}', 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        contact = doc.add_paragraph('{{EMAIL}} | {{PHONE}} | {{ADDRESS}}')
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if '{{LINKEDIN}}':
            linkedin = doc.add_paragraph('{{LINKEDIN}}')
            linkedin.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Summary
        doc.add_heading('Professional Summary', level=1)
        doc.add_paragraph('{{SUMMARY}}')
        
        # Experience
        doc.add_heading('Work Experience', level=1)
        doc.add_paragraph('{{EXPERIENCE_SECTION}}')
        
        # Education
        doc.add_heading('Education', level=1)
        doc.add_paragraph('{{EDUCATION_SECTION}}')
        
        # Skills
        doc.add_heading('Technical Skills', level=1)
        doc.add_paragraph('{{SKILLS_SECTION}}')
        
        # Save
        template_path = os.path.join(self.templates_dir, 'professional.docx')
        doc.save(template_path)
        logger.info(f"Created professional template: {template_path}")
    
    def _create_modern_template(self):
        """Create a modern template with placeholders"""
        doc = Document()
        
        # Header with modern styling
        header = doc.add_heading('{{FULL_NAME}}', 0)
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        doc.add_paragraph('{{EMAIL}} • {{PHONE}}')
        doc.add_paragraph('{{ADDRESS}}')
        doc.add_paragraph('{{LINKEDIN}}')
        
        # About section
        doc.add_heading('About', level=1)
        doc.add_paragraph('{{SUMMARY}}')
        
        # Experience with modern layout
        doc.add_heading('Experience', level=1)
        doc.add_paragraph('{{EXPERIENCE_SECTION}}')
        
        # Education
        doc.add_heading('Education', level=1)
        doc.add_paragraph('{{EDUCATION_SECTION}}')
        
        # Skills
        doc.add_heading('Core Competencies', level=1)
        doc.add_paragraph('{{SKILLS_SECTION}}')
        
        # Save
        template_path = os.path.join(self.templates_dir, 'modern.docx')
        doc.save(template_path)
        logger.info(f"Created modern template: {template_path}")
